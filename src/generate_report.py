from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = PROJECT_ROOT / "data_processed"
DOCS_DIR = PROJECT_ROOT / "docs"
SOURCE_URL = "https://databank.worldbank.org/source/international-debt-statistics"


def money_billions(value: float) -> str:
    return f"${value:,.2f}B"


def pct(value: float) -> str:
    return f"{value:,.2f}%"


def load_inputs() -> dict:
    summary_path = DATA_DIR / "data_quality_summary.json"
    with summary_path.open("r", encoding="utf-8") as f:
        quality = json.load(f)

    main_year = quality["summary"]["main_year"]
    return {
        "quality": quality,
        "main_year": main_year,
        "top_countries": pd.read_csv(DATA_DIR / f"top_countries_total_external_debt_{main_year}.csv"),
        "indicator_totals": pd.read_csv(DATA_DIR / f"indicator_totals_{main_year}.csv"),
        "region_totals": pd.read_csv(DATA_DIR / f"region_total_external_debt_{main_year}.csv"),
        "trend": pd.read_csv(DATA_DIR / "global_total_external_debt_trend.csv"),
    }


def insight_text(inputs: dict) -> list[str]:
    summary = inputs["quality"]["summary"]
    top = inputs["top_countries"]
    regions = inputs["region_totals"]
    trend = inputs["trend"]

    total = summary["total_external_debt_billions_usd"]
    top_10_share = top.head(10)["value_billions_usd"].sum() / total * 100
    top_region = regions.iloc[0]
    top_region_share = top_region["value_billions_usd"] / total * 100
    first_year = int(trend.iloc[0]["year"])
    last_year = int(trend.iloc[-1]["year"])
    change = (
        (trend.iloc[-1]["value_billions_usd"] - trend.iloc[0]["value_billions_usd"])
        / trend.iloc[0]["value_billions_usd"]
        * 100
    )

    return [
        (
            f"In {summary['main_year']}, total external debt across the analysed countries was "
            f"{money_billions(total)}."
        ),
        (
            f"{summary['top_country']} had the highest total external debt at "
            f"{money_billions(summary['top_country_debt_billions_usd'])}."
        ),
        (
            f"The top 10 countries accounted for {pct(top_10_share)} of total external debt, "
            "showing a high concentration among a small group of borrowers."
        ),
        (
            f"{top_region['region']} was the largest regional contributor with "
            f"{money_billions(top_region['value_billions_usd'])}, or {pct(top_region_share)} "
            "of the analysed total."
        ),
        (
            f"From {first_year} to {last_year}, total external debt changed by {pct(change)}, "
            "indicating the longer-term direction of global debt exposure."
        ),
    ]


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        header[i].text = str(value)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def build_docx(inputs: dict, output_path: Path) -> None:
    quality = inputs["quality"]
    summary = quality["summary"]
    main_year = inputs["main_year"]
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading("International Debt Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Python, SQL, and Streamlit analytics project")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Objective", level=1)
    document.add_paragraph(
        "This project converts raw World Bank International Debt Statistics data into a "
        "clean analytics dataset, a normalized SQL model, a set of analytical SQL queries, "
        "and an interactive dashboard for country-wise and indicator-wise debt analysis."
    )

    document.add_heading("Dataset", level=1)
    document.add_paragraph(f"Source: World Bank International Debt Statistics ({SOURCE_URL}).")
    document.add_paragraph(
        f"The processed dataset covers {quality['year_start']} to {quality['year_end']} and "
        "keeps World-counterpart monetary debt indicators measured in current US dollars."
    )

    document.add_heading("Data Preparation", level=1)
    for item in [
        "Loaded the raw CSV export using Python and Pandas.",
        "Removed duplicate records and reshaped year columns into a long-format table.",
        "Filtered relevant monetary debt indicators and retained non-null observations.",
        "Joined country and indicator metadata to support richer analysis.",
        "Created normalized country, indicator, and debt fact tables for MySQL.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Key Metrics", level=1)
    add_docx_table(
        document,
        [
            ["Metric", "Value"],
            ["Main reporting year", str(main_year)],
            ["Processed rows", f"{quality['cleaned_rows']:,}"],
            ["Countries", f"{summary['country_count']:,}"],
            ["Indicators", f"{summary['indicator_count']:,}"],
            ["Total external debt", money_billions(summary["total_external_debt_billions_usd"])],
            ["Highest-debt country", summary["top_country"]],
            ["Lowest-debt country", summary["lowest_country"]],
        ],
    )

    document.add_heading("Top Countries by Total External Debt", level=1)
    top_rows = [["Rank", "Country", "Region", "Debt"]]
    for index, row in inputs["top_countries"].head(10).reset_index(drop=True).iterrows():
        top_rows.append(
            [
                str(index + 1),
                row["country_name"],
                row["region"],
                money_billions(row["value_billions_usd"]),
            ]
        )
    add_docx_table(document, top_rows)

    document.add_heading("Regional Debt Distribution", level=1)
    region_rows = [["Region", "Debt"]]
    for _, row in inputs["region_totals"].iterrows():
        region_rows.append([row["region"], money_billions(row["value_billions_usd"])])
    add_docx_table(document, region_rows)

    document.add_heading("Top Debt Indicators", level=1)
    indicator_rows = [["Indicator", "Measure Type", "Value"]]
    for _, row in inputs["indicator_totals"].head(10).iterrows():
        indicator_rows.append(
            [row["series_name"], row["measure_type"], money_billions(row["value_billions_usd"])]
        )
    add_docx_table(document, indicator_rows)

    document.add_heading("Insights", level=1)
    for item in insight_text(inputs):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Deliverables", level=1)
    for item in [
        "Cleaned CSV dataset and normalized CSV tables.",
        "Complete MySQL schema with primary and foreign key relationships.",
        "Thirty SQL analytical queries covering basic, intermediate, and advanced analysis.",
        "Streamlit dashboard for interactive exploration.",
        "Final documentation in DOCX and PDF formats.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Limitations", level=1)
    document.add_paragraph(
        "The analysis uses current-US-dollar debt indicators and does not adjust for inflation, "
        "exchange-rate effects, or country population size. Some indicators are related financial "
        "flows rather than debt stock, so total external debt uses the dedicated indicator "
        "DT.DOD.DECT.CD."
    )

    document.save(output_path)


def pdf_table(rows: list[list[str]], widths: list[float]) -> Table:
    table = Table(rows, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C6CF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FA")]),
            ]
        )
    )
    return table


def build_pdf(inputs: dict, output_path: Path) -> None:
    quality = inputs["quality"]
    summary = quality["summary"]
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []

    story.append(Paragraph("International Debt Analysis", styles["Title"]))
    story.append(Paragraph("Python, SQL, and Streamlit analytics project", styles["Normal"]))
    story.append(Spacer(1, 0.18 * inch))

    story.append(Paragraph("Objective", styles["Heading1"]))
    story.append(
        Paragraph(
            "This project converts raw World Bank International Debt Statistics data into a "
            "clean analytics dataset, normalized SQL tables, analytical SQL queries, and an "
            "interactive dashboard.",
            styles["BodyText"],
        )
    )

    story.append(Paragraph("Dataset", styles["Heading1"]))
    story.append(Paragraph(f"Source: World Bank International Debt Statistics - {SOURCE_URL}", styles["BodyText"]))
    story.append(
        Paragraph(
            f"Processed range: {quality['year_start']} to {quality['year_end']}. "
            "The dataset keeps World-counterpart monetary debt indicators measured in current US dollars.",
            styles["BodyText"],
        )
    )

    story.append(Paragraph("Key Metrics", styles["Heading1"]))
    story.append(
        pdf_table(
            [
                ["Metric", "Value"],
                ["Main year", str(inputs["main_year"])],
                ["Processed rows", f"{quality['cleaned_rows']:,}"],
                ["Countries", f"{summary['country_count']:,}"],
                ["Indicators", f"{summary['indicator_count']:,}"],
                ["Total external debt", money_billions(summary["total_external_debt_billions_usd"])],
                ["Highest-debt country", summary["top_country"]],
                ["Lowest-debt country", summary["lowest_country"]],
            ],
            [2.5 * inch, 4 * inch],
        )
    )

    story.append(Paragraph("Insights", styles["Heading1"]))
    for item in insight_text(inputs):
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(PageBreak())

    story.append(Paragraph("Top Countries by Total External Debt", styles["Heading1"]))
    top_rows = [["Rank", "Country", "Region", "Debt"]]
    for index, row in inputs["top_countries"].head(10).reset_index(drop=True).iterrows():
        top_rows.append(
            [
                str(index + 1),
                row["country_name"],
                row["region"],
                money_billions(row["value_billions_usd"]),
            ]
        )
    story.append(pdf_table(top_rows, [0.45 * inch, 1.7 * inch, 2.55 * inch, 1.15 * inch]))

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Regional Debt Distribution", styles["Heading1"]))
    region_rows = [["Region", "Debt"]]
    for _, row in inputs["region_totals"].iterrows():
        region_rows.append([row["region"], money_billions(row["value_billions_usd"])])
    story.append(pdf_table(region_rows, [4.0 * inch, 1.7 * inch]))

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Top Debt Indicators", styles["Heading1"]))
    indicator_rows = [["Indicator", "Measure", "Value"]]
    for _, row in inputs["indicator_totals"].head(8).iterrows():
        indicator_rows.append(
            [row["series_name"], row["measure_type"], money_billions(row["value_billions_usd"])]
        )
    story.append(pdf_table(indicator_rows, [3.8 * inch, 1.35 * inch, 0.9 * inch]))

    story.append(Paragraph("Limitations", styles["Heading1"]))
    story.append(
        Paragraph(
            "The analysis uses current-US-dollar values and does not adjust for inflation, "
            "exchange-rate effects, or population. Total external debt is calculated using "
            "indicator DT.DOD.DECT.CD.",
            styles["BodyText"],
        )
    )

    doc.build(story)


def build_markdown(inputs: dict, output_path: Path) -> None:
    quality = inputs["quality"]
    summary = quality["summary"]
    lines = [
        "# International Debt Analysis",
        "",
        "## Objective",
        "Build a complete Python, SQL, and dashboard analytics project using World Bank International Debt Statistics.",
        "",
        "## Dataset",
        f"Source: {SOURCE_URL}",
        f"Processed years: {quality['year_start']} to {quality['year_end']}",
        "",
        "## Key Metrics",
        f"- Main year: {inputs['main_year']}",
        f"- Processed rows: {quality['cleaned_rows']:,}",
        f"- Countries: {summary['country_count']:,}",
        f"- Indicators: {summary['indicator_count']:,}",
        f"- Total external debt: {money_billions(summary['total_external_debt_billions_usd'])}",
        f"- Highest-debt country: {summary['top_country']}",
        f"- Lowest-debt country: {summary['lowest_country']}",
        "",
        "## Insights",
    ]
    lines.extend(f"- {item}" for item in insight_text(inputs))
    lines.extend(
        [
            "",
            "## Deliverables",
            "- Cleaned CSV dataset and normalized tables.",
            "- MySQL schema and 30 analytical SQL queries.",
            "- Streamlit dashboard.",
            "- Final DOCX and PDF report.",
            "",
        ]
    )
    output_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    inputs = load_inputs()
    build_docx(inputs, DOCS_DIR / "International_Debt_Analysis_Report.docx")
    build_pdf(inputs, DOCS_DIR / "International_Debt_Analysis_Report.pdf")
    build_markdown(inputs, DOCS_DIR / "EDA_Report.md")
    print("Generated DOCX, PDF, and Markdown reports.")


if __name__ == "__main__":
    main()
